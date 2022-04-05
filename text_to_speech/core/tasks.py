import base64
from io import BytesIO

import requests
from celery import shared_task
from django.conf import settings
from docx import Document
from docx.shared import Inches
from storages.backends.s3boto3 import S3Boto3Storage

from .utils import milliseconds_to_hours


@shared_task
def transcribe(file):

    headers = {
        'authorization': settings.ASSEMBLYAI_TOKEN,
        'content-type': 'application/json',
    }

    file = file.encode()
    file = base64.b64decode(file)

    # Step 1: Upload file
    response = requests.post(
        'https://api.assemblyai.com/v2/upload', headers=headers, data=file
    )
    upload_url = response.json()['upload_url']

    # Step 2: Submit file for transcription
    endpoint = 'https://api.assemblyai.com/v2/transcript'
    json_data = {'audio_url': upload_url, 'speaker_labels': True}
    response = requests.post(endpoint, json=json_data, headers=headers)
    polling_endpoint = endpoint + '/' + response.json()['id']

    # Step 3: Poll Assembly Ai
    status = 'submited'
    while status not in ('completed', 'error'):
        response = requests.get(polling_endpoint, headers=headers)
        status = response.json()['status']

    if status == 'error':
        raise Exception('Transcription failed')

    # Step 4: Create word document and return the url
    utterances = response.json()['utterances']
    document = Document()

    for utterance in utterances:
        start = milliseconds_to_hours(utterance['start'])
        try:
            document.add_picture(
                'staticfiles/core/images/Judiciary-Logo.jpg',
                width=Inches(1.0),
                height=Inches(1.0),
            )
        except:
            pass
        document.add_paragraph(f'{utterance["speaker"]} at {start}')
        document.add_paragraph(f'{utterance["text"]}')
        document.add_paragraph()

    with BytesIO() as fileobj:
        document.save(fileobj)
        fileobj.seek(0)

        storage = S3Boto3Storage()
        file_name = 'transcription.docx'
        filename = storage.save(file_name, fileobj)
        file_url = storage.url(filename)

    return file_url
